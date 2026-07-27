import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

doc = docx.Document()

# BaşlıkStili
title = doc.add_heading(
    "🚀 Çevrimdışı (Offline) Yerel RAG Asistanı Proje Dökümantasyonu", level=0
)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Amac
p = doc.add_paragraph()
p.add_run("Proje Amacı: ").bold = True
p.add_run(
    "İnternet bağlantısına ve bulut servislerine ihtiyaç duymadan, tamamen yerel (local) "
    "bilgisayar kaynaklarıyla akademik dökümanları (PDF) okuyan, veritabanında saklayan ve döküman "
    "hakkında sorulan sorulara doğru cevaplar veren yapay zeka asistanı geliştirmek."
)

# PHASE 1
doc.add_heading(
    "📌 PHASE 1: WEEK 1 – Altyapı, Yerel Model ve Sistem Kurulumu", level=1
)
doc.add_heading("💡 Teorik Açıklama", level=2)
doc.add_paragraph(
    "1. Bulut Yerin Yerel Yapay Zeka (Local AI): Verilerimizin dışarıya çıkmaması ve tamamen ücretsiz çalışması için Microsoft'un yerel yapay zeka çalıştırma altyapısı olan Foundry Local runtime'ını kullandık.\n"
    "2. Model Seçimi (qwen2.5-0.5b): Bilgisayarımızın işlemci ve RAM kaynaklarını yormayacak, hızlı yanıt veren 0.5 milyar parametreli hafif Dil Modelini (LLM) yerel hafızaya yükledik.\n"
    "3. Takma Ad (Alias) Mantığı: Windows Terminalinde uzun dosya yolları yazmak yerine Set-Alias komutuyla bilgisayara kısayol tanımlamayı öğrendik."
)

# PHASE 2
doc.add_heading(
    "📌 PHASE 2: WEEK 2 – Döküman Okuma ve Metin Dilimleme (Chunking)", level=1
)
doc.add_heading("💡 Teorik Açıklama", level=2)
doc.add_paragraph(
    "1. Bağlam Sınırı (Token Limit) Problemi: Yapay zeka modellerine 10-20 sayfalık makalelerin tamamını tek seferde gönderirsek model 'Token Limit Exceeded' hatası verir.\n"
    "2. Metin Dilimleme (Chunking): Bu sorunu çözmek için LangChain kütüphanesini kullanarak PDF belgesini mantıksal paragraflara böldük.\n"
    "3. Çakışma Payı (Overlap): Parçaları bölerken cümlelerin ortadan kesilip anlam kaybı yaşamaması için her parçanın son 40-60 karakterini bir sonraki parçanın başına dahil ettik."
)

# PHASE 3
doc.add_heading(
    "📌 PHASE 3: WEEK 3 – Vektörleştirme (Embedding), SQLite ve Akıllı Arama",
    level=1,
)
doc.add_heading("💡 Teorik Açıklama", level=2)
doc.add_paragraph(
    "1. Metin Vektörleştirme (Embedding): BAAI/bge-small-en-v1.5 modelini kullanarak her metin parçasını 384 boyutlu matematiksel sayı dizilerine (vektörlere) dönüştürdük.\n"
    "2. Gürültü Filtreleme: Makale sonundaki Kaynakça ve DOI linkleri aramalarda yanlış eşleşmeye sebep olduğu için otomatik temizledik.\n"
    "3. SQLite Veritabanı (rag_memory.db): Vektörleştirilen parçaları tek bir veritabanı dosyasında kalıcı olarak sakladık.\n"
    "4. Kosinüs Benzerliği Matematiği: Soru ile veritabanındaki parçalar arasındaki açıyı Kosinüs Benzerliği formülü ile hesapladık."
)

# Sunum Tüyoları
doc.add_heading(
    "🎯 3 Altın Cümle", level=1
)
doc.add_paragraph(
    "• Week 1 Mantığı: 'Burada buluta tek bir kuruş ödemeden Microsoft Foundry Local ile doğrudan kendi bilgisayarımızın işlemcisinde çalışan yerel yapay zekayı ayağa kaldırdık.'\n"
    "• Week 2 Mantığı: 'Aşırı uzun dökümanları doğrudan modele verirsek yapay zekanın kafası karışıyor ve hafızası yetmiyor. Bu yüzden dökümanı 400'er karakterlik küçük anlamlı dilimlere bölüyoruz.'\n"
    "• Week 3 Mantığı: 'Kelimeleri aramıyoruz, anlamları arıyoruz! BGE modeliyle her parçayı matematiksel koordinatlara dönüştürüp SQLite veritabanına gömdük. Soru sorduğumuzda veritabanı en yakın anlama sahip 2 parçayı %80 isabetle çat diye önümüze getiriyor.'"
)

doc.save("RAG_Proje_Dokumantasyonu.docx")
print("✅ Word dosyası 'RAG_Proje_Dokumantasyonu.docx' adıyla oluşturuldu!")